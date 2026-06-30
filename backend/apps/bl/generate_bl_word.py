"""
Génération du document Word (.docx) du Bon de Livraison (BL), équivalent au
PDF : en-tête entreprise, client, tableau des déchets, chauffeur, signature.
"""
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .generate_bl import _recuperateur_info, _destinataire_info


def generate_bl_docx(data: dict) -> bytes:
    rec  = _recuperateur_info(data)
    dest = _destinataire_info(data)

    doc = Document()

    if rec['logo_path']:
        try:
            doc.add_picture(rec['logo_path'], width=Cm(2.2))
        except Exception:
            pass

    nom_p = doc.add_paragraph()
    nom_run = nom_p.add_run(rec['nom'].upper())
    nom_run.bold = True
    nom_run.italic = True
    nom_run.font.size = Pt(16)

    if rec['agrement_num']:
        doc.add_paragraph(f"Agrément N° {rec['agrement_num']} du {rec['agrement_date']}")
    adresse_ligne = ' '.join(filter(None, [rec['adresse'], rec['code_postal']]))
    if adresse_ligne:
        doc.add_paragraph(adresse_ligne)
    doc.add_paragraph(f"RC {rec['rc']}\tNIF {rec['nif']}")
    doc.add_paragraph(f"NA {rec['na']}\tNIS {rec['nis']}")

    doc.add_paragraph()
    lieu_date = doc.add_paragraph(f"{rec['commune']} le : {data.get('date_livraison','')}")
    lieu_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph(f"Nom de Client : {dest['nom']}")
    doc.add_paragraph(f"Adresse : {dest['adresse']}")

    doc.add_paragraph()
    titre = doc.add_heading('Bon de livraison', level=2)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    lignes = data.get('lignes') or []
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['N°', 'Description (Nature des déchets)', 'Quantités', 'Unités', 'Stockage']):
        hdr[i].text = h
    for i, l in enumerate(lignes, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(l.get('description', ''))
        row[2].text = str(l.get('quantite', ''))
        row[3].text = str(l.get('unite', 'KG'))
        row[4].text = str(l.get('stockage', ''))

    doc.add_paragraph()
    doc.add_paragraph(f"Nom de chauffeur : {data.get('chauffeur_nom','')}")
    doc.add_paragraph(f"Immatriculation de camion : {data.get('camion_immatriculation','')}")

    doc.add_paragraph()
    doc.add_paragraph()
    gerant = doc.add_paragraph('Le Gérant')
    gerant.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if rec['responsable']:
        resp = doc.add_paragraph(rec['responsable'])
        resp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
