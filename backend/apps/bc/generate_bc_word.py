"""
Génération du document Word (.docx) du Bon de Commande (BC) :
en-tête récupérateur, client, tableau avec prix et totaux, signature.
"""
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .generate_bc import _recuperateur_info, _calc_totaux


def generate_bc_docx(data: dict) -> bytes:
    rec     = _recuperateur_info(data)
    lignes  = data.get('lignes') or []
    tva_pct = float(data.get('tva_pct') or 19)

    doc = Document()

    # ── Marges ──────────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.left_margin  = Cm(2)
        section.right_margin = Cm(2)

    # ── Logo ────────────────────────────────────────────────────────────────────
    if rec['logo_path']:
        try:
            doc.add_picture(rec['logo_path'], width=Cm(2.2))
        except Exception:
            pass

    # ── En-tête récupérateur ────────────────────────────────────────────────────
    nom_p = doc.add_paragraph()
    run   = nom_p.add_run(rec['nom'].upper())
    run.bold   = True
    run.italic = True
    run.font.size  = Pt(16)
    run.font.color.rgb = RGBColor(0x3B, 0x6D, 0x11)

    if rec['agrement_num']:
        doc.add_paragraph(f"Agrément N° {rec['agrement_num']} du {rec['agrement_date']}")
    adresse = ' '.join(filter(None, [rec['adresse'], rec['code_postal']]))
    if adresse:
        doc.add_paragraph(adresse)
    doc.add_paragraph(f"RC {rec['rc']}\tNIF {rec['nif']}")
    doc.add_paragraph(f"NA {rec['na']}\tNIS {rec['nis']}")

    # ── Date / lieu ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p_date = doc.add_paragraph(f"{rec['commune']} le : {data.get('date_commande', '')}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ── Client ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph(f"Nome de Client : {data.get('client_nom', '')}")
    doc.add_paragraph(f"Adresse : {data.get('client_adresse', '')}")

    # ── Titre ───────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    titre = doc.add_heading('Bon de commande', level=2)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Tableau des déchets ──────────────────────────────────────────────────────
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['N°', 'Description (Nature des déchets)', 'Quantités', 'Unités', 'Prix unitaires', 'Total HT']):
        hdr[i].text = h

    total_ht_sum = 0.0
    for i, l in enumerate(lignes, start=1):
        try:
            qte = float(l.get('quantite') or 0)
            pu  = float(l.get('prix_unitaire') or 0)
            ht  = qte * pu
        except (TypeError, ValueError):
            ht = 0.0
        total_ht_sum += ht
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(l.get('description', ''))
        row[2].text = str(l.get('quantite', ''))
        row[3].text = str(l.get('unite', 'KG'))
        row[4].text = f"{pu:,.2f} DZ".replace(',', ' ')
        row[5].text = f"{ht:,.2f} DZ".replace(',', ' ')

    # ── Récapitulatif ───────────────────────────────────────────────────────────
    total_ht, tva, total_ttc = _calc_totaux(lignes, tva_pct)
    doc.add_paragraph()
    recap = doc.add_table(rows=3, cols=2)
    recap.style = 'Table Grid'
    recap.rows[0].cells[0].text = 'Total HT'
    recap.rows[0].cells[1].text = f"{total_ht:,.2f} DZ".replace(',', ' ')
    recap.rows[1].cells[0].text = f'TVA ({tva_pct:.0f}%)'
    recap.rows[1].cells[1].text = f"{tva:,.2f} DZ".replace(',', ' ')
    recap.rows[2].cells[0].text = 'Total TTC'
    recap.rows[2].cells[1].text = f"{total_ttc:,.2f} DZ".replace(',', ' ')
    for row in recap.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Right-align the whole table
    tbl_pr = recap._tbl
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    tbl_pr.tblPr.append(jc)

    # ── Arrêté ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    arrete = doc.add_paragraph(
        f"Arrêter la présente facture en toutes taxes comprises a la somme de : "
        f"{total_ttc:,.2f} DZ".replace(',', ' ')
    )
    arrete.runs[0].italic = True

    # ── Signature ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    gerant = doc.add_paragraph('Le Gérant')
    gerant.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if rec['responsable']:
        resp = doc.add_paragraph(rec['responsable'])
        resp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
